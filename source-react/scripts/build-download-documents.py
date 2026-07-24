from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "downloads"
OUT.mkdir(parents=True, exist_ok=True)

GREEN = "6B8E23"
LIME = "B9F232"
DARK = "151814"
MUTED = "687066"
PALE = "F5F8F1"
LINE = "DDE4D8"
ORANGE = "EA6B28"

YEARS = ["2024", "2025", "2026F", "2027F", "2028F"]
FINANCIALS = [
    ["매출액 (십억원)", "17,941", "17,099", "17,941", "18,337", "18,712"],
    ["영업이익 (십억원)", "1,823", "1,073", "1,935", "2,153", "2,435"],
    ["EPS (원)", "5,810", "1,901", "6,347", "7,205", "8,424"],
    ["P/E (배)", "9.5", "28.1", "13.2", "11.6", "10.0"],
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.13

    for name, size, color in [("Title", 28, DARK), ("Heading 1", 17, DARK), ("Heading 2", 12.5, GREEN)]:
        style = styles[name]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(14 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(7)

    header = section.header.paragraphs[0]
    header.text = "REFLO  |  EQUITY RESEARCH"
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(GREEN)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "REFLO Research Workspace  ·  SK텔레콤 2Q26 Earnings Review"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(9)
    run = kicker.add_run("017670 · 유무선통신  |  2026. 7. 17")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(ORANGE)

    doc.add_heading("SK텔레콤 2Q26 Earnings Review", 0)
    subtitle = doc.add_paragraph("호실적과 업종 내 AIDC 확장의 결합")
    subtitle.style = styles["Subtitle"]
    subtitle.runs[0].font.name = "맑은 고딕"
    subtitle.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    subtitle.runs[0].font.size = Pt(17)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    kpi = doc.add_table(rows=2, cols=4)
    kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi.autofit = False
    kpi_data = [
        ["투자의견", "목표주가", "현재주가", "상승여력"],
        ["매수", "120,000원", "83,900원", "43.0%"],
    ]
    for r, row in enumerate(kpi.rows):
        for c, cell in enumerate(row.cells):
            cell.width = Inches(1.55)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, PALE if r == 0 else "FFFFFF")
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(kpi_data[r][c])
            run.font.name = "맑은 고딕"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            run.font.size = Pt(9 if r == 0 else 14)
            run.font.bold = r == 1
            run.font.color.rgb = RGBColor.from_string(MUTED if r == 0 else (GREEN if c in (0, 3) else DARK))

    doc.add_heading("Investment Summary", level=1)
    p = doc.add_paragraph()
    p.add_run("2분기 실적 상회와 AIDC 사업의 가시성 개선을 함께 반영해 ").bold = False
    strong = p.add_run("투자의견 매수, 목표주가 120,000원")
    strong.bold = True
    strong.font.color.rgb = RGBColor.from_string(GREEN)
    p.add_run("을 유지합니다. 현재 주가 대비 상승여력은 43.0%입니다.")

    doc.add_heading("핵심 투자 포인트", level=2)
    points = [
        ("본업 이익 체력 회복", "무선 서비스의 비용 효율화와 유선 가입자 증가를 바탕으로 2026년 영업이익은 1조 9,350억원까지 정상화될 전망입니다."),
        ("AIDC 가치 재평가", "현재 137MW인 데이터센터 수전용량은 2027년 187MW로 확대되고, 울산 AIDC 가동과 추가 부지 확보가 중장기 성장 경로를 구체화합니다."),
        ("주주환원 가시성", "2026년 예상 DPS 3,660원과 배당수익률 4.4%는 대규모 AI 투자 구간에서도 하방 경직성을 제공합니다."),
    ]
    for idx, (title, body) in enumerate(points, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(f"{title}  ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(DARK)
        p.add_run(body)

    doc.add_heading("실적 전망", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["구분", *YEARS]
    for c, text in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.width = Cm(3.6 if c == 0 else 2.25)
        set_cell_shading(cell, DARK)
        set_cell_border(cell, DARK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    set_repeat_table_header(table.rows[0])
    for row_data in FINANCIALS:
        cells = table.add_row().cells
        for c, text in enumerate(row_data):
            cell = cells[c]
            cell.width = Cm(3.6 if c == 0 else 2.25)
            set_cell_shading(cell, PALE if c == 0 else "FFFFFF")
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.bold = c == 0 or (c == 3)
            run.font.size = Pt(9)
            if c == 3:
                run.font.color.rgb = RGBColor.from_string(GREEN)

    doc.add_heading("AIDC 전망", level=1)
    doc.add_paragraph("2027년부터 수전용량과 가동률 상승이 실적에 반영될 전망입니다. 2035년 15GW 확장 계획은 중장기 성장의 선택지를 넓히지만, 준공 일정과 전력 조달비는 지속적으로 확인해야 합니다.")

    doc.add_heading("밸류에이션", level=1)
    value = doc.add_table(rows=2, cols=3)
    value.alignment = WD_TABLE_ALIGNMENT.CENTER
    value_data = [["Forward EPS", "Target PER", "계산 목표주가"], ["12,430원", "14.2배", "176,506원"]]
    for r, row in enumerate(value.rows):
        for c, cell in enumerate(row.cells):
            set_cell_shading(cell, PALE if r == 0 else ("EFF5E7" if c == 2 else "FFFFFF"))
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value_data[r][c])
            run.font.size = Pt(9 if r == 0 else 15)
            run.bold = r == 1
            run.font.color.rgb = RGBColor.from_string(GREEN if c == 2 else DARK)

    doc.add_heading("주요 리스크", level=1)
    for title, body in [
        ("AIDC 투자 집행과 가동 지연", "전력 인입, 인허가, 고객 유치가 계획보다 늦어질 경우 초기 투자비 부담이 먼저 반영될 수 있습니다."),
        ("무선 가입자 성장 둔화", "5G 보급률이 성숙기에 진입하면서 가입자 순증과 ARPU 개선 속도가 예상보다 낮을 수 있습니다."),
        ("주주환원 여력 축소", "대규모 데이터센터 투자와 차입금 증가는 배당 정상화 속도를 제한할 가능성이 있습니다."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{title}: ").bold = True
        p.add_run(body)

    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(14)
    run = source.add_run("자료: 회사 공시, 기업 IR, REFLO Research Workspace. 본 문서는 예시 리서치 산출물입니다.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    path = OUT / "SK_Telecom_2Q26_Report.docx"
    doc.save(path)
    return path


def add_pdf():
    regular = r"C:\Windows\Fonts\malgun.ttf"
    bold = r"C:\Windows\Fonts\malgunbd.ttf"
    pdfmetrics.registerFont(TTFont("Malgun", regular))
    pdfmetrics.registerFont(TTFont("MalgunBold", bold))

    path = OUT / "SK_Telecom_2Q26_Report.pdf"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=17 * mm,
        bottomMargin=16 * mm,
        title="SK텔레콤 2Q26 Earnings Review",
        author="REFLO Research Workspace",
    )
    styles = getSampleStyleSheet()
    base = ParagraphStyle("BaseKR", fontName="Malgun", fontSize=9.3, leading=14, textColor=colors.HexColor(f"#{DARK}"), spaceAfter=6)
    title = ParagraphStyle("TitleKR", parent=base, fontName="MalgunBold", fontSize=24, leading=31, spaceAfter=5)
    subtitle = ParagraphStyle("SubtitleKR", parent=base, fontSize=15, leading=20, textColor=colors.HexColor(f"#{MUTED}"), spaceAfter=13)
    h1 = ParagraphStyle("H1KR", parent=base, fontName="MalgunBold", fontSize=15, leading=20, spaceBefore=12, spaceAfter=8)
    h2 = ParagraphStyle("H2KR", parent=base, fontName="MalgunBold", fontSize=11, leading=15, textColor=colors.HexColor(f"#{GREEN}"), spaceBefore=5, spaceAfter=5)
    kicker = ParagraphStyle("KickerKR", parent=base, fontName="MalgunBold", fontSize=8.5, leading=12, textColor=colors.HexColor(f"#{ORANGE}"), spaceAfter=8)
    note = ParagraphStyle("NoteKR", parent=base, fontSize=7.8, leading=11, textColor=colors.HexColor(f"#{MUTED}"))
    center = ParagraphStyle("CenterKR", parent=base, alignment=TA_CENTER)

    story = [
        Paragraph("REFLO · EQUITY RESEARCH", kicker),
        Paragraph("SK텔레콤 2Q26 Earnings Review", title),
        Paragraph("호실적과 업종 내 AIDC 확장의 결합", subtitle),
    ]
    kpi_data = [
        ["투자의견", "목표주가", "현재주가", "상승여력"],
        ["매수", "120,000원", "83,900원", "43.0%"],
    ]
    kpi_table = Table(kpi_data, colWidths=[42 * mm] * 4, rowHeights=[8 * mm, 12 * mm])
    kpi_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Malgun"),
        ("FONTNAME", (0, 1), (-1, 1), "MalgunBold"),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("FONTSIZE", (0, 1), (-1, 1), 13),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(f"#{MUTED}")),
        ("TEXTCOLOR", (0, 1), (-1, 1), colors.HexColor(f"#{DARK}")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{PALE}")),
        ("BACKGROUND", (0, 1), (-1, 1), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{LINE}")),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.extend([kpi_table, Spacer(1, 6 * mm), Paragraph("Investment Summary", h1)])
    story.append(Paragraph("2분기 실적 상회와 AIDC 사업의 가시성 개선을 함께 반영해 <b>투자의견 매수, 목표주가 120,000원</b>을 유지합니다. 현재 주가 대비 상승여력은 43.0%입니다.", base))
    story.append(Paragraph("핵심 투자 포인트", h2))
    for idx, (point_title, body) in enumerate([
        ("본업 이익 체력 회복", "무선 서비스 비용 효율화와 유선 가입자 증가로 2026년 영업이익은 1조 9,350억원까지 정상화될 전망입니다."),
        ("AIDC 가치 재평가", "데이터센터 수전용량은 2027년 187MW로 확대되고, 울산 AIDC 가동과 추가 부지가 중장기 성장 경로를 구체화합니다."),
        ("주주환원 가시성", "2026년 예상 DPS 3,660원과 배당수익률 4.4%는 대규모 AI 투자 구간의 하방 경직성을 제공합니다."),
    ], 1):
        story.append(Paragraph(f"<b>{idx:02d}  {point_title}</b><br/>{body}", base))

    story.extend([Spacer(1, 3 * mm), Paragraph("실적 전망", h1)])
    fin_data = [["구분", *YEARS], *FINANCIALS]
    fin_table = Table(fin_data, colWidths=[37 * mm] + [26 * mm] * 5, repeatRows=1)
    fin_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "MalgunBold"),
        ("FONTNAME", (0, 1), (-1, -1), "Malgun"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{DARK}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BACKGROUND", (0, 1), (0, -1), colors.HexColor(f"#{PALE}")),
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor(f"#{LINE}")),
        ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TEXTCOLOR", (3, 1), (3, -1), colors.HexColor(f"#{GREEN}")),
    ]))
    story.append(fin_table)
    story.extend([
        Paragraph("AIDC 전망", h1),
        Paragraph("2027년부터 수전용량과 가동률 상승이 실적에 반영될 전망입니다. 2035년 15GW 확장 계획은 성장의 선택지를 넓히지만 준공 일정과 전력 조달비는 지속적으로 확인해야 합니다.", base),
        Paragraph("밸류에이션", h1),
    ])
    val_data = [["Forward EPS", "×", "Target PER", "=", "계산 목표주가"], ["12,430원", "×", "14.2배", "=", "176,506원"]]
    val_table = Table(val_data, colWidths=[38 * mm, 12 * mm, 38 * mm, 12 * mm, 54 * mm], rowHeights=[8 * mm, 14 * mm])
    val_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Malgun"),
        ("FONTNAME", (0, 1), (-1, 1), "MalgunBold"),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("FONTSIZE", (0, 1), (-1, 1), 14),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(f"#{MUTED}")),
        ("TEXTCOLOR", (4, 1), (4, 1), colors.HexColor(f"#{GREEN}")),
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{PALE}")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{LINE}")),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(val_table)
    story.extend([Paragraph("주요 리스크", h1)])
    for idx, (risk_title, body) in enumerate([
        ("AIDC 투자 집행과 가동 지연", "전력 인입, 인허가, 고객 유치가 계획보다 늦어질 경우 초기 투자비 부담이 먼저 반영될 수 있습니다."),
        ("무선 가입자 성장 둔화", "5G 보급률이 성숙기에 진입하면서 가입자 순증과 ARPU 개선 속도가 예상보다 낮을 수 있습니다."),
        ("주주환원 여력 축소", "대규모 데이터센터 투자와 차입금 증가는 배당 정상화 속도를 제한할 가능성이 있습니다."),
    ], 1):
        story.append(Paragraph(f"<b>{idx:02d}  {risk_title}</b><br/>{body}", base))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("자료: 회사 공시, 기업 IR, REFLO Research Workspace. 본 문서는 예시 리서치 산출물입니다.", note))

    def page_deco(canvas, pdf_doc):
        canvas.saveState()
        width, height = A4
        canvas.setStrokeColor(colors.HexColor(f"#{LIME}"))
        canvas.setLineWidth(1.1)
        canvas.line(18 * mm, height - 10 * mm, width - 18 * mm, height - 10 * mm)
        canvas.setFont("Malgun", 7)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawString(18 * mm, 8 * mm, "REFLO Research Workspace")
        canvas.drawRightString(width - 18 * mm, 8 * mm, f"{pdf_doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=page_deco, onLaterPages=page_deco)
    return path


if __name__ == "__main__":
    print(add_docx())
    print(add_pdf())
